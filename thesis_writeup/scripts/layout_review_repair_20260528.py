from __future__ import annotations

import shutil
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


DOCX = Path(r"C:\Users\lby\Downloads\v2\thesis_writeup\dissertation_final.docx")
BACKUP_DIR = Path(r"C:\Users\lby\Downloads\v2\thesis_writeup\archive_old_versions_20260528_layout_review")
BACKUP = BACKUP_DIR / "dissertation_final_before_layout_review.docx"


def has_page_break_before(paragraph) -> bool:
    p_pr = paragraph._p.pPr
    if p_pr is None:
        return False
    return p_pr.find(qn("w:pageBreakBefore")) is not None


def set_page_break_before(paragraph) -> bool:
    if has_page_break_before(paragraph):
        return False
    p_pr = paragraph._p.get_or_add_pPr()
    p_pr.append(OxmlElement("w:pageBreakBefore"))
    return True


def remove_trailing_empty_paragraphs(doc: Document) -> int:
    removed = 0
    body = doc.element.body
    for paragraph in reversed(doc.paragraphs):
        text = paragraph.text.strip()
        xml = paragraph._p.xml
        if text:
            break
        if "<w:sectPr" in xml:
            break
        body.remove(paragraph._p)
        removed += 1
    return removed


def update_static_front_matter(doc: Document) -> int:
    page_numbers = {
        "1 Introduction": "1",
        "1.1 Background": "1",
        "1.2 Problem Statement": "1",
        "1.3 Aim and Objectives": "2",
        "1.4 Research Questions": "2",
        "1.5 Outcomes": "3",
        "1.6 Thesis Structure": "3",
        "2 Literature Review": "4",
        "3 Methodology": "13",
        "4 LIAR In-Domain Experiments": "21",
        "5 Error Analysis and Further Experiments": "25",
        "6 Cross-Dataset Evaluation": "29",
        "7 Discussion": "38",
        "8 Conclusion": "42",
        "References": "45",
        "Appendices": "48",
        "3.1 Overall experimental pipeline for cross-dataset fake news detection": "14",
        "3.2 Binary label mapping and class distribution across LIAR and FakeNewsNet": "16",
        "3.3 Five-seed transformer training and evaluation protocol": "18",
        "4.1 LIAR in-domain performance comparison across local baselines": "22",
        "4.2 REAL and FAKE recall trade-off on the LIAR test set": "23",
        "5.1 Qualitative taxonomy of recurring LIAR error patterns": "26",
        "6.1 Macro-F1 on LIAR and FakeNewsNet combined titles": "31",
        "6.2 Class-level recall on FakeNewsNet combined titles": "32",
        "6.3 Majority baseline comparison on FakeNewsNet combined titles": "32",
        "6.4 Prediction bias under strict transfer": "35",
        "A.1 Reproducible experiment file structure and output artifacts": "51",
        "1.1 Research questions": "2",
        "3.1 Dataset splits and label distributions for LIAR and FakeNewsNet": "14",
        "3.2 Binary label mapping": "15",
        "3.3 Baseline models": "17",
        "3.4 Main transformer hyperparameter settings": "17",
        "4.1 Main LIAR results": "21",
        "4.2 Class-level recall for local models with saved recall metrics": "22",
        "5.1 Representative LIAR error cases discussed in the main text": "27",
        "6.1 LIAR in-domain comparison before transfer": "30",
        "6.2 LIAR-to-FakeNewsNet title transfer results": "30",
        "6.3 Majority-class (always-REAL) reference baseline on FakeNewsNet minimal titles": "30",
        "A.1 Classification metric summaries used in the dissertation": "48",
        "B.1 Available confusion matrices and matrix sources": "49",
        "C.1 Reproducible hyperparameter settings": "49",
        "D.1 Representative LIAR error-analysis examples": "50",
    }
    changed = 0
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if "\t" not in text:
            continue
        title = text.rsplit("\t", 1)[0].strip()
        page = page_numbers.get(title)
        if page is None:
            continue
        new_text = f"{title} \t{page}"
        if paragraph.text != new_text:
            paragraph.text = new_text
            changed += 1
    return changed


def update_word_fields(docx_path: Path) -> None:
    import pythoncom
    import win32com.client

    pythoncom.CoInitialize()
    word = win32com.client.DispatchEx("Word.Application")
    word.Visible = False
    word.DisplayAlerts = 0
    try:
        doc = word.Documents.OpenNoRepairDialog(str(docx_path), False, False, False)
        doc.Repaginate()
        for story in doc.StoryRanges:
            current = story
            while current is not None:
                current.Fields.Update()
                try:
                    current = current.NextStoryRange
                except Exception:
                    current = None
        for field in doc.Fields:
            field.Update()
        for toc in doc.TablesOfContents:
            toc.Update()
            toc.UpdatePageNumbers()
        for tof in doc.TablesOfFigures:
            tof.Update()
            tof.UpdatePageNumbers()
        doc.Repaginate()
        doc.Save()
        doc.Close(False)
    finally:
        word.Quit()
        pythoncom.CoUninitialize()


def main() -> None:
    BACKUP_DIR.mkdir(parents=True, exist_ok=True)
    if not BACKUP.exists():
        shutil.copy2(DOCX, BACKUP)

    doc = Document(DOCX)
    targets = {
        "3.8 Reproducibility",
        "5.7 Summary against RQ3",
        "7.4 Future Work",
    }
    breaks_added = []
    for paragraph in doc.paragraphs:
        if paragraph.text.strip() in targets:
            if set_page_break_before(paragraph):
                breaks_added.append(paragraph.text.strip())

    trailing_removed = remove_trailing_empty_paragraphs(doc)
    static_front_matter_updated = update_static_front_matter(doc)
    doc.save(DOCX)
    update_word_fields(DOCX)
    print(
        {
            "docx": str(DOCX),
            "backup": str(BACKUP),
            "breaks_added": breaks_added,
            "trailing_empty_paragraphs_removed": trailing_removed,
            "static_front_matter_updated": static_front_matter_updated,
        }
    )


if __name__ == "__main__":
    main()
