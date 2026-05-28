from __future__ import annotations

import shutil
import os
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt


ROOT = Path(r"C:\Users\lby\Downloads\v2")
THESIS = ROOT / "thesis_writeup"
DOCX = Path(os.environ.get("DISSERTATION_DOCX", THESIS / "dissertation_final.docx"))
BACKUP = THESIS / "archive_old_versions_20260527_figure_repair" / "dissertation_final_before_table_repair.docx"


def find_paragraph(doc: Document, startswith: str):
    for idx, paragraph in enumerate(doc.paragraphs):
        if " ".join(paragraph.text.split()).startswith(startswith):
            return idx, paragraph
    raise ValueError(f"paragraph not found: {startswith}")


def paragraph_has_page_break(paragraph) -> bool:
    return bool(paragraph._p.xpath(".//w:br[@w:type='page']"))


def insert_page_break_before(paragraph) -> None:
    previous = paragraph._p.getprevious()
    if previous is not None and previous.xpath(".//w:br[@w:type='page']"):
        return
    new_p = OxmlElement("w:p")
    run = OxmlElement("w:r")
    br = OxmlElement("w:br")
    br.set(qn("w:type"), "page")
    run.append(br)
    new_p.append(run)
    paragraph._p.addprevious(new_p)


def set_cell_shading(cell, fill: str = "FFFFFF") -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    for shd in tc_pr.findall(qn("w:shd")):
        tc_pr.remove(shd)
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_width(cell, width_dxa: int) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_cell_margins(cell, margin_dxa: int = 80) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    mar = tc_pr.find(qn("w:tcMar"))
    if mar is None:
        mar = OxmlElement("w:tcMar")
        tc_pr.append(mar)
    for side in ("top", "left", "bottom", "right"):
        elem = mar.find(qn(f"w:{side}"))
        if elem is None:
            elem = OxmlElement(f"w:{side}")
            mar.append(elem)
        elem.set(qn("w:w"), str(margin_dxa))
        elem.set(qn("w:type"), "dxa")


def set_table_borders(table) -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        elem = borders.find(qn(f"w:{edge}"))
        if elem is None:
            elem = OxmlElement(f"w:{edge}")
            borders.append(elem)
        elem.set(qn("w:val"), "single")
        elem.set(qn("w:sz"), "6")
        elem.set(qn("w:space"), "0")
        elem.set(qn("w:color"), "000000")


def prevent_row_splits(table) -> None:
    for row in table.rows:
        tr_pr = row._tr.get_or_add_trPr()
        if tr_pr.find(qn("w:cantSplit")) is None:
            tr_pr.append(OxmlElement("w:cantSplit"))


def set_table_fixed_widths(table, widths: list[int]) -> None:
    tbl_pr = table._tbl.tblPr
    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            if idx < len(widths):
                set_cell_width(cell, widths[idx])


def normalize_table(table, font_size: float = 9.5, widths: list[int] | None = None) -> None:
    table.style = "Table Grid"
    set_table_borders(table)
    prevent_row_splits(table)
    if widths:
        set_table_fixed_widths(table, widths)
    for r_idx, row in enumerate(table.rows):
        for c_idx, cell in enumerate(row.cells):
            set_cell_shading(cell, "FFFFFF")
            set_cell_margins(cell, 90)
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_before = Pt(0)
                paragraph.paragraph_format.space_after = Pt(0)
                paragraph.paragraph_format.line_spacing = 1.0
                paragraph.paragraph_format.keep_together = True
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    run.font.name = "Times New Roman"
                    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
                    run.font.size = Pt(font_size)
                    if r_idx == 0:
                        run.bold = True


def main() -> None:
    if not BACKUP.exists():
        BACKUP.parent.mkdir(parents=True, exist_ok=True)
        shutil.copy2(DOCX, BACKUP)

    doc = Document(str(DOCX))

    _, table_61_caption = find_paragraph(doc, "Table 6.1.")
    insert_page_break_before(table_61_caption)
    table_61_caption.paragraph_format.keep_with_next = True

    _, table_42_caption = find_paragraph(doc, "Table 4.2.")
    insert_page_break_before(table_42_caption)
    table_42_caption.paragraph_format.keep_with_next = True

    _, table_64_caption = find_paragraph(doc, "Table 6.4.")
    table_64_caption.paragraph_format.keep_with_next = True

    # Current table indices are stable in dissertation_final.docx.
    table_61 = doc.tables[6]
    table_64 = doc.tables[12]

    normalize_table(table_61, font_size=10.5, widths=[2500, 900, 1550, 1550, 1550, 1550])
    normalize_table(table_64, font_size=8.7, widths=[1300, 1600, 1050, 1200, 1200, 1200, 1810])

    # Keep the row labels readable in the wide integrated table.
    for row in table_64.rows[1:]:
        for c_idx in (0, 1, 6):
            for paragraph in row.cells[c_idx].paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.save(str(DOCX))
    print(DOCX)


if __name__ == "__main__":
    main()
