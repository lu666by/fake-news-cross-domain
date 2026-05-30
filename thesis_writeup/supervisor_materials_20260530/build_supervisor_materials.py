from __future__ import annotations

import re
import unicodedata
from pathlib import Path

import pandas as pd
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[2]
OUT_DIR = Path(__file__).resolve().parent
DATE = "2026-05-30"


def ascii_text(text: str) -> str:
    replacements = {
        "\u2014": " - ",
        "\u2013": "-",
        "\u2018": "'",
        "\u2019": "'",
        "\u201c": '"',
        "\u201d": '"',
        "\u00a0": " ",
    }
    for old, new in replacements.items():
        text = text.replace(old, new)
    text = unicodedata.normalize("NFKD", text)
    text = text.encode("ascii", "ignore").decode("ascii")
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r" *\n *", "\n", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


def chapter_file_by_size(size: int) -> Path:
    files = list((ROOT / "thesis_writeup" / "dissertation_planning").rglob("*.txt"))
    matches = [p for p in files if p.stat().st_size == size]
    if len(matches) != 1:
        raise RuntimeError(f"Expected one chapter file of size {size}, found {matches}")
    return matches[0]


def replace_block(text: str, start: str, end: str, replacement: str) -> str:
    i = text.index(start)
    j = text.index(end, i) + len(end)
    return text[:i] + replacement.strip() + text[j:]


def normalize_chapter_headings(text: str) -> str:
    lines = text.splitlines()
    out: list[str] = []
    for n, line in enumerate(lines):
        stripped = line.strip()
        if not stripped:
            out.append("")
            continue
        if stripped.startswith("Chapter "):
            out.append(f"# {stripped}")
        elif n == 1 and lines[0].startswith("Chapter "):
            out.append(f"## {stripped}")
        elif re.match(r"^\d+\.\d+\.\d+\s+", stripped):
            out.append(f"### {stripped}")
        elif re.match(r"^\d+\.\d+\s+", stripped):
            out.append(f"## {stripped}")
        elif stripped.startswith("Table "):
            out.append(f"**{stripped}**")
        else:
            out.append(stripped)
    return "\n".join(out)


def format_metric(value: float | str) -> str:
    if value == "" or pd.isna(value):
        return ""
    return f"{float(value):.4f}"


def build_results_summary() -> str:
    rows = [
        [
            "Titles-only direct transfer",
            "LIAR train -> held-out FakeNewsNet title test",
            "5 seeds",
            "0.2725 +/- 0.0179",
            "0.2364 +/- 0.0257",
            "0.0377 +/- 0.0268",
            "0.9842 +/- 0.0110",
            "Consistently weak; not a seed-42 accident.",
        ],
        [
            "Intermediate fine-tuning 10%",
            "LIAR train -> stratified 10% FNN target-train titles -> held-out test",
            "5 seeds",
            "0.8083 +/- 0.0068",
            "0.7035 +/- 0.0298",
            "0.9305 +/- 0.0197",
            "0.4379 +/- 0.0791",
            "Stable uplift over direct transfer; wording should say stratified, not 1:1 balanced.",
        ],
        [
            "Intermediate fine-tuning 20%",
            "LIAR train -> stratified 20% FNN target-train titles -> held-out test",
            "5 seeds",
            "0.8243 +/- 0.0049",
            "0.7463 +/- 0.0060",
            "0.9167 +/- 0.0138",
            "0.5444 +/- 0.0292",
            "Best absolute 5-seed target-fraction result so far; compare with 10% for efficiency.",
        ],
        [
            "1000-row reasoning-atom pilot",
            "LLM atoms with logistic decision layer",
            "pilot",
            "0.5300 / 0.6000",
            "0.4405 / 0.6000",
            "0.1300 / 0.5900",
            "0.9300 / 0.6100",
            "Current DeepSeek V4 Flash reasoning-atom setup shows limited performance in this pilot.",
        ],
    ]
    header = [
        "Experiment",
        "Protocol",
        "Evidence",
        "Accuracy",
        "Macro-F1",
        "REAL recall",
        "FAKE recall",
        "Reading",
    ]
    lines = [
        "| " + " | ".join(header) + " |",
        "| " + " | ".join(["---"] * len(header)) + " |",
    ]
    for row in rows:
        lines.append("| " + " | ".join(row) + " |")
    return "\n".join(lines)


TERMS_TABLE = """| Term | Meaning in this dissertation |
| --- | --- |
| target fraction | The proportion of FakeNewsNet target-train titles used after the LIAR source model has been trained. In the current rerun, 10% means a stratified 10% sample from each target label, not a 1:1 balanced subset. |
| target training | The FakeNewsNet target-domain training subset used for adaptation. It is separate from validation and held-out test data. |
| intermediate fine-tuning | A two-stage training protocol: first fine-tune on LIAR, then continue fine-tuning on a small target-domain FakeNewsNet training subset before final held-out target testing. |
| held-out test | The final FakeNewsNet test split. It is not used for model training, checkpoint selection, threshold tuning, or explanation-driven adjustment. |
| seed | The random seed controlling model initialization, training order, and target-domain sampling/splitting. Multiple seeds are used to test whether a result is stable rather than accidental. |
| pilot | A small exploratory run used to test whether an idea is promising. A pilot can motivate discussion, but it should not be written as a strong conclusion without reruns or additional samples. |"""


CHAPTER3_TERMS_PROSE = """3.4 Experimental Terminology for Transfer and Fine-Tuning

This dissertation uses several experimental terms in a narrow and reproducible way. The target fraction is the proportion of FakeNewsNet target-train titles used after the LIAR-trained source model has already been trained. For example, the 10% condition uses 1,484 target-train titles per seed. The current code samples this subset stratified by label, so "10%" should be read as a stratified target fraction rather than a 1:1 balanced training subset.

Target training refers only to the FakeNewsNet training subset used for adaptation. It is not the held-out target test set. Intermediate fine-tuning means that the model is first fine-tuned on LIAR and then further fine-tuned on this small FakeNewsNet target-training subset. The held-out test is the final FakeNewsNet test split; it is kept separate from training, validation, checkpoint selection, and threshold tuning.

A seed is the random seed controlling initialization, training order, and target-domain sampling. Reporting several seeds is important because fine-tuned transformer models can vary across random runs. A pilot is a small exploratory experiment, such as the 1000-row reasoning-atom run, that can guide discussion but should not be presented as a strong conclusion unless it is repeated or enlarged."""


def prepare_chapters() -> tuple[str, str]:
    chapter2 = chapter_file_by_size(24598).read_text(encoding="utf-8")
    chapter3 = chapter_file_by_size(4790).read_text(encoding="utf-8")

    chapter2 = replace_block(
        chapter2,
        "Table 2.1. Summary of prior literature on fake news classification and dataset benchmarks.",
        "Papageorgiou [8]      BERT, Traditional, LLMs           LIAR, FakeNewsNet       ~65.0% - 68.0% Accuracy         Evaluated multiple datasets independently; did not run strict zero-shot transfer.",
        """Table 2.1. Summary of prior literature on fake news classification and dataset benchmarks.

| Study | Model(s) Used | Dataset(s) | Reported Metric(s) | Key Findings and Limitations |
| --- | --- | --- | --- | --- |
| Wang [1] | SVM, Logistic Regression | LIAR (6-class) | ~27.7% Accuracy | Baseline paper; metadata improves accuracy slightly but text-only remains low. |
| Rashkin et al. [4] | MaxEnt, Naive Bayes, LSTM | LIAR (binary) | ~62.0% Macro-F1 | Found stylistic and sentiment patterns; highly specific to political domain. |
| Perez-Rosas [5] | SVM with linguistic and lexical features | Custom, FakeNewsNet | ~70.0%-80.0% Accuracy | Strong in-domain performance but 15-20 point drops out of domain. |
| Devlin et al. [6] | BERT-Base / Large | GLUE benchmarks | - | Established bidirectional representation fine-tuning; no transfer tested. |
| Liu et al. [7] | RoBERTa-Base | GLUE, SQuAD | - | Optimised BERT pre-training; transfer under distribution shift not evaluated. |
| Papageorgiou [8] | BERT, traditional models, LLMs | LIAR, FakeNewsNet | ~65.0%-68.0% Accuracy | Evaluated multiple datasets independently; did not run strict zero-shot transfer. |""",
    )

    old_overview = """This dissertation uses a staged experimental design. The first stage builds reliable in-domain baselines on the LIAR dataset. The second stage analyses whether class-weighted training changes model behaviour, especially for the FAKE class. The third stage measures whether models trained on LIAR statements remain useful when applied to FakeNewsNet titles without any adaptation, against an always-REAL class-prior baseline.

The design keeps the project controlled and reproducible. All main models use the same binary label mapping, the same LIAR train, validation, and test splits, and the same statement field as input. This makes the comparison between TF-IDF, BERT, weighted BERT, and weighted RoBERTa controlled rather than changing several variables at the same time.

The cross-dataset setting is deliberately strict. No FakeNewsNet examples are used for training, validation, threshold selection, or checkpoint selection. This means the target-domain result measures transfer from LIAR rather than adaptation to FakeNewsNet, establishing a rigorous zero-shot benchmark for transferability limits."""
    new_overview = """This dissertation uses a staged experimental design. The first stage builds reliable in-domain baselines on the LIAR dataset. The second stage analyses whether class-weighted training changes model behaviour, especially for the FAKE class. The third stage measures strict titles-only direct transfer, where models trained on LIAR statements are evaluated on FakeNewsNet titles without target-domain adaptation. The fourth stage introduces limited intermediate fine-tuning on a small FakeNewsNet target-training subset to test whether a modest amount of target-domain supervision can reduce the direct-transfer failure.

The design keeps the project controlled and reproducible. All main models use the same binary label mapping, the same LIAR train, validation, and test splits, and the same statement field as input. This makes the comparison between TF-IDF, BERT, weighted BERT, and weighted RoBERTa controlled rather than changing several variables at the same time.

The direct-transfer setting is deliberately strict: no FakeNewsNet examples are used for source-model training, threshold selection, or checkpoint selection. The intermediate fine-tuning setting is separate from this zero-shot baseline. It uses a clearly defined FakeNewsNet target-training subset after LIAR training, while keeping the held-out FakeNewsNet test split untouched for final evaluation."""
    chapter3 = chapter3.replace(old_overview, new_overview)

    chapter3 = replace_block(
        chapter3,
        "Table 3.2. Dataset splits and label distributions for LIAR and FakeNewsNet.",
        "                     Combined Titles        17,441 (75.19%)     5,755 (24.81%)      23,196",
        """Table 3.2. Dataset splits and label distributions for LIAR and FakeNewsNet.

| Dataset | Split | REAL Count (%) | FAKE Count (%) | Total Rows |
| --- | --- | ---: | ---: | ---: |
| LIAR (Source) | Train | 5,752 (56.17%) | 4,488 (43.83%) | 10,240 |
| LIAR (Source) | Validation | 668 (52.02%) | 616 (47.98%) | 1,284 |
| LIAR (Source) | Test | 714 (56.35%) | 553 (43.65%) | 1,267 |
| FakeNewsNet (Target) | PolitiFact Titles | 624 (59.09%) | 432 (40.91%) | 1,056 |
| FakeNewsNet (Target) | GossipCop Titles | 16,817 (75.96%) | 5,323 (24.04%) | 22,140 |
| FakeNewsNet (Target) | Combined Titles | 17,441 (75.19%) | 5,755 (24.81%) | 23,196 |""",
    )

    chapter3 = replace_block(
        chapter3,
        "Table 3.1. Binary label mapping.",
        "FAKE              barely-true, false, pants-on-fire",
        """Table 3.1. Binary label mapping.

| Binary label | Original LIAR labels |
| --- | --- |
| REAL | true, mostly-true, half-true |
| FAKE | barely-true, false, pants-on-fire |""",
    )
    chapter3 = chapter3.replace("3.4 Model Configuration", CHAPTER3_TERMS_PROSE + "\n\n3.5 Model Configuration")

    return normalize_chapter_headings(ascii_text(chapter2)), normalize_chapter_headings(ascii_text(chapter3))


def build_markdown(chapter2: str, chapter3: str) -> str:
    return f"""# Supervisor Reading Version: Chapters 2 and 3 with Updated Results Summary

Date: {DATE}

This version is intended for supervisor reading before the next meeting. It prioritizes Chapter 2 and Chapter 3, plus a concise summary of the updated rerun evidence. Chapter 6 has been updated internally, but it should not be the main material sent for close reading yet.

## What changed this week

- Re-ran the titles-only direct transfer baseline on five seeds.
- Re-ran the 10% and 20% intermediate fine-tuning settings on five seeds.
- Updated the integrated results table with the new five-seed means.
- Repositioned the 1000-row reasoning-atom experiment as a current-setup pilot.
- Added body-text explanations for target fraction, target training, intermediate fine-tuning, held-out test, seed, and pilot.

## Updated Results Summary

{build_results_summary()}

The main result is now easier to state: titles-only direct transfer is consistently weak, while 10% and 20% intermediate fine-tuning both give stable improvements across seeds. The 10% result is the more data-efficient setting; the 20% result is currently the strongest absolute five-seed setting. The 1000-row atom result should be described only as a current DeepSeek V4 Flash reasoning-atom setup with limited performance in this pilot.

## Terminology Notes for Reader

{TERMS_TABLE}

{chapter2}

{chapter3}
"""


def build_chapter6_internal() -> str:
    integrated = (ROOT / "results" / "integrated_experiment_summary" / "integrated_main_results_table.md").read_text(
        encoding="utf-8"
    )
    integrated = ascii_text(integrated)
    return f"""# Chapter 6 Internal Rerun Update

Date: {DATE}

Status: internal draft. Do not make this the main supervisor reading package yet. The supervisor-facing material should lead with Chapter 2, Chapter 3, and the short updated results summary.

## Internal Claims That Are Now Safer

- The titles-only direct transfer baseline is consistently weak across five seeds. The 5-seed mean Macro-F1 is 0.2364, and the model remains strongly FAKE-biased on FakeNewsNet titles.
- The 10% intermediate fine-tuning setting is stable across five seeds. Mean Macro-F1 is 0.7035, compared with 0.2364 for the paired direct-transfer baseline.
- The 20% intermediate fine-tuning setting is also stable across five seeds and is currently the strongest absolute target-fraction result, with mean Macro-F1 0.7463 and mean Accuracy 0.8243.
- The 10% subset must be described as stratified target-title training, not as a 1:1 balanced subset.
- The completed 20% rerun now supports a five-seed claim. The remaining judgement is whether the thesis should emphasize 10% for data efficiency or 20% for best absolute performance.
- The 1000-row TELLER-like atom experiment should remain a current-setup pilot. It can support motivation and future work, but it should not be framed as a strong empirical conclusion from one random sample.

## Suggested Chapter 6 Wording

The five-seed rerun confirms that the poor direct-transfer result is not an artefact of seed 42. Across seeds 42, 52, 62, 72, and 82, the LIAR-trained weighted RoBERTa model achieves only 0.2364 mean Macro-F1 on the held-out FakeNewsNet title test set and predicts FAKE for most target examples. This establishes a stable source-only transfer failure rather than an isolated bad run.

By contrast, 10% intermediate fine-tuning gives a stable improvement over the direct baseline. The model is first fine-tuned on LIAR and then further fine-tuned on a stratified 10% FakeNewsNet target-training title subset. Across five seeds, this setting reaches 0.7035 mean Macro-F1 and 0.8083 mean accuracy on the same held-out target test split. The improvement is therefore large and repeatable, although FAKE recall remains lower than REAL recall.

The 20% target-fraction rerun confirms that the larger target subset is also stable across the same five seeds. It reaches 0.7463 mean Macro-F1 and 0.8243 mean accuracy, making it the best absolute five-seed target-fraction result so far. The thesis framing should therefore distinguish between the 10% setting as the efficient adaptation result and the 20% setting as the strongest absolute result.

## Updated Integrated Table

{integrated}
"""


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_width(cell, width_dxa: int) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.first_child_found_in("w:tcW")
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths: list[int]) -> None:
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), "9360")
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.first_child_found_in("w:tblInd")
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")
    tbl_grid = table._tbl.tblGrid
    if tbl_grid is None:
        tbl_grid = OxmlElement("w:tblGrid")
        table._tbl.insert(0, tbl_grid)
    for child in list(tbl_grid):
        tbl_grid.remove(child)
    for width in widths:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(width))
        tbl_grid.append(grid_col)
    for row in table.rows:
        for i, cell in enumerate(row.cells):
            set_cell_width(cell, widths[min(i, len(widths) - 1)])


def add_paragraph(doc: Document, text: str, style: str | None = None):
    para = doc.add_paragraph(style=style)
    para.add_run(text)
    return para


def add_markdown_table(doc: Document, lines: list[str]) -> None:
    rows: list[list[str]] = []
    for line in lines:
        cells = [c.strip() for c in line.strip().strip("|").split("|")]
        if all(set(c) <= {":", "-"} for c in cells):
            continue
        rows.append(cells)
    if not rows:
        return
    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    table.style = "Table Grid"
    widths = [max(900, int(9360 / len(rows[0])))] * len(rows[0])
    if len(rows[0]) == 2:
        widths = [1900, 7460]
    elif len(rows[0]) == 5:
        widths = [1400, 1700, 1500, 1500, 3260]
    elif len(rows[0]) == 8:
        widths = [1400, 2200, 900, 900, 900, 900, 900, 1260]
    set_table_geometry(table, widths)
    for r_idx, row in enumerate(rows):
        for c_idx, cell_text in enumerate(row):
            cell = table.cell(r_idx, c_idx)
            cell.text = ""
            p = cell.paragraphs[0]
            run = p.add_run(cell_text)
            run.font.name = "Calibri"
            run.font.size = Pt(8 if len(rows[0]) >= 5 else 9)
            if r_idx == 0:
                run.bold = True
                set_cell_shading(cell, "F2F4F7")
    doc.add_paragraph()


def build_docx(markdown: str, out_path: Path) -> None:
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    styles["Normal"].font.name = "Calibri"
    styles["Normal"].font.size = Pt(11)
    styles["Normal"].paragraph_format.space_after = Pt(6)
    styles["Normal"].paragraph_format.line_spacing = 1.10
    for name, size, color, before, after in [
        ("Heading 1", 16, "2E74B5", 16, 8),
        ("Heading 2", 13, "2E74B5", 12, 6),
        ("Heading 3", 12, "1F4D78", 8, 4),
    ]:
        style = styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)

    pending_table: list[str] = []
    title_done = False

    def flush_table() -> None:
        nonlocal pending_table
        if pending_table:
            add_markdown_table(doc, pending_table)
            pending_table = []

    for raw_line in markdown.splitlines():
        line = raw_line.rstrip()
        if not line:
            flush_table()
            continue
        if line.startswith("|"):
            pending_table.append(line)
            continue
        flush_table()
        if line.startswith("# "):
            text = line[2:].strip()
            if not title_done:
                para = doc.add_paragraph()
                para.alignment = WD_ALIGN_PARAGRAPH.LEFT
                run = para.add_run(text)
                run.font.name = "Calibri"
                run.font.size = Pt(20)
                run.font.bold = True
                run.font.color.rgb = RGBColor.from_string("0B2545")
                title_done = True
            else:
                add_paragraph(doc, text, "Heading 1")
        elif line.startswith("## "):
            add_paragraph(doc, line[3:].strip(), "Heading 2")
        elif line.startswith("### "):
            add_paragraph(doc, line[4:].strip(), "Heading 3")
        elif line.startswith("- "):
            para = doc.add_paragraph(style="List Bullet")
            para.add_run(line[2:].strip())
        elif line.startswith("**") and line.endswith("**"):
            if "Table 3.1." in line:
                doc.add_page_break()
            para = doc.add_paragraph()
            run = para.add_run(line.strip("*"))
            run.bold = True
        else:
            add_paragraph(doc, line)
    flush_table()

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    footer.add_run("Supervisor reading version, " + DATE)
    doc.save(out_path)


def main() -> None:
    chapter2, chapter3 = prepare_chapters()
    supervisor_md = build_markdown(chapter2, chapter3)
    internal_ch6 = build_chapter6_internal()

    md_path = OUT_DIR / f"supervisor_ch2_ch3_results_{DATE}.md"
    ch6_path = OUT_DIR / f"chapter6_internal_rerun_update_{DATE}.md"
    docx_path = OUT_DIR / f"supervisor_ch2_ch3_results_{DATE}.docx"

    md_path.write_text(supervisor_md, encoding="utf-8")
    ch6_path.write_text(internal_ch6, encoding="utf-8")
    build_docx(supervisor_md, docx_path)

    print(md_path)
    print(ch6_path)
    print(docx_path)


if __name__ == "__main__":
    main()
